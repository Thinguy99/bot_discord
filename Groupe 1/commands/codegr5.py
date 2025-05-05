import requests
import json
import os
from docx import Document
from typing import Dict, List, Optional, Union, Tuple, Any

class ModuleGroupe5:
    """
    Module pour l'analyse de CV et la génération de lettres de motivation
    Développé par le Groupe 5 : Aymane AIBICHI, Zineb MANAR, Ali BOUGUERRA, Nawel ARIF, Nhung Nguyen.
    
    Cette classe permet de :
    - Comparer un CV avec une fiche de poste pour évaluer la pertinence (>70%)
    - Générer une lettre de motivation personnalisée basée sur le CV et l'offre
    """
    
    def __init__(self, api_key: str, output_dir: str = "lettres"):
        """
        Initialise le module groupe 5
        
        Args:
            api_key: Clé API Google Gemini
            output_dir: Répertoire où seront sauvegardées les lettres générées
        """
        self.api_key = api_key
        self.output_dir = output_dir
        self.url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-pro:generateContent"
        self.headers = {"Content-Type": "application/json"}
        
        # Créer le répertoire de sortie s'il n'existe pas
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
    
    # --- Fonctions pour formatter les données du CV ---
    def _formatter_formation(self, formation: List[Dict]) -> str:
        """Formate la section formation du CV pour le prompt"""
        return "\n".join(
            f"- {f['titre']} – {f['etablissement']} ({f['periode']})\n  " +
            "\n  ".join(f['details']) for f in formation
        )

    def _formatter_experience(self, experience: List[Dict]) -> str:
        """Formate la section expérience du CV pour le prompt"""
        return "\n".join(
            f"- {e['titre']} – {e['entreprise']}, {e['lieu']} ({e['periode']})\n  " +
            "\n  ".join(e['details']) for e in experience
        )
    
    # --- Génération des prompts ---
    def generer_prompt_pertinence(self, cv_dict: Dict, offre_dict: Dict) -> str:
        """
        Génère le prompt pour évaluer la pertinence du CV par rapport à l'offre
        
        Args:
            cv_dict: Dictionnaire contenant les informations du CV
            offre_dict: Dictionnaire contenant les informations de l'offre
            
        Returns:
            Prompt formaté pour l'API Gemini
        """
        formations = self._formatter_formation(cv_dict["formation"])
        experiences = self._formatter_experience(cv_dict["experience"])
        competences = "\n- ".join(cv_dict["competences_techniques"])
        soft_skills = "\n- ".join(cv_dict["soft_skills"])
        langues = "\n- ".join(cv_dict["langues"])
        certifications = "\n- ".join(cv_dict["certifications"])

        cv_txt = f"""Nom : {cv_dict['prenom_nom']}
Email : {cv_dict['email']}
Téléphone : {cv_dict['telephone']}

Formation :
{formations}

Expériences :
{experiences}

Compétences techniques :
- {competences}

Soft skills :
- {soft_skills}

Langues :
- {langues}

Certifications :
- {certifications}
"""

        offre_txt = f"""Titre : {offre_dict['titre']}
Entreprise : {offre_dict['entreprise']}
Lieu : {offre_dict['lieu']}
Contrat : {offre_dict['type_contrat']}

Description de l'entreprise :
{offre_dict['description_entreprise']}

Missions :
{offre_dict['missions']}

Profil recherché :
{offre_dict['profil_recherche']}
"""

        return f"""
Tu es un expert RH.

Voici un CV et une offre d'emploi. Réponds uniquement par "oui" si le profil correspond à plus de 70 % à l'offre, sinon réponds "non". Ne donne aucune explication.

--- CV ---
{cv_txt}

--- Offre ---
{offre_txt}
"""

    def generer_prompt_lettre(self, cv_dict: Dict, offre_dict: Dict, infos_perso: Optional[Dict] = None) -> str:
        """
        Génère le prompt pour la création d'une lettre de motivation
        
        Args:
            cv_dict: Dictionnaire contenant les informations du CV
            offre_dict: Dictionnaire contenant les informations de l'offre
            infos_perso: Dictionnaire contenant des informations personnelles supplémentaires
            
        Returns:
            Prompt formaté pour l'API Gemini
        """
        if infos_perso is None:
            infos_perso = {"motivation": "", "lien_entreprise": "", "contraintes": ""}

        formations = self._formatter_formation(cv_dict["formation"])
        experiences = self._formatter_experience(cv_dict["experience"])
        competences = ", ".join(cv_dict["competences_techniques"])
        soft_skills = ", ".join(cv_dict["soft_skills"])
        langues = ", ".join(cv_dict["langues"])
        certifications = ", ".join(cv_dict["certifications"])

        perso_txt = ""
        if infos_perso["motivation"]:
            perso_txt += f"\nMotivation personnelle : {infos_perso['motivation']}"
        if infos_perso["lien_entreprise"]:
            perso_txt += f"\nLien particulier avec l'entreprise ou le secteur : {infos_perso['lien_entreprise']}"
        if infos_perso["contraintes"]:
            perso_txt += f"\nInformations supplémentaires : {infos_perso['contraintes']}"

        return f"""
Tu es un expert RH et spécialiste de la rédaction de lettres de motivation professionnelles. Rédige une lettre complète, prête à être envoyée, en t'appuyant sur le CV du candidat et l'offre d'emploi ci-dessous.

🎯 Objectif :
Fournir une lettre claire, convaincante, personnalisée, sans faute ni besoin de correction, dans un style fluide, professionnel et humain.

✅ La lettre doit impérativement :
- Tenir sur une page (Word A4) avec un style direct et efficace.
- Suivre ce plan structuré :
    1. Présentation brève du candidat et de son parcours
    2. Motivation sincère et cohérente pour le poste
    3. Mise en lien entre l'entreprise/l'offre et les valeurs du candidat
    4. Mise en avant ciblée des compétences, expériences ou cours suivis correspondant aux missions
    5. Remerciements, disponibilité pour un entretien, et formule de politesse

✍️ Style :
- Zéro faute d'orthographe ou de grammaire.
- Chaque phrase commence par une majuscule.
- Aucune formule générique ni tournure artificielle.
- Le ton doit être confiant, positif, professionnel et chaleureux.
- Ne propose aucun espace à compléter : tout doit être finalisé.

📎 Contexte fourni :

--- CV ---
Nom : {cv_dict['prenom_nom']}
Email : {cv_dict['email']}
Téléphone : {cv_dict['telephone']}
LinkedIn : {cv_dict.get('linkedin', 'Non spécifié')}
GitHub : {cv_dict.get('github', 'Non spécifié')}

Formation :
{formations}

Expériences :
{experiences}

Compétences techniques : {competences}
Compétences comportementales (soft skills) : {soft_skills}
Langues : {langues}
Certifications : {certifications}

--- Offre ---
Titre : {offre_dict['titre']}
Entreprise : {offre_dict['entreprise']}
Lieu : {offre_dict['lieu']}
Type de contrat : {offre_dict['type_contrat']}
À propos de l'entreprise :
{offre_dict['description_entreprise']}

Missions proposées :
{offre_dict['missions']}

Profil recherché :
{offre_dict['profil_recherche']}

--- Informations complémentaires du candidat ---
{perso_txt if perso_txt else "Aucune information supplémentaire fournie."}
"""

    # --- Interaction avec l'API Gemini ---
    def interroger_gemini(self, prompt: str) -> Optional[str]:
        """
        Envoie un prompt à l'API Gemini et récupère la réponse
        
        Args:
            prompt: Le prompt à envoyer à l'API
            
        Returns:
            La réponse de l'API ou None en cas d'erreur
        """
        data = {
            "contents": [
                {
                    "role": "user",
                    "parts": [{"text": prompt}]
                }
            ]
        }
        params = {"key": self.api_key}
        
        try:
            response = requests.post(self.url, headers=self.headers, params=params, json=data)
            if response.status_code == 200:
                resultat = response.json()
                return resultat['candidates'][0]['content']['parts'][0]['text'].strip()
            else:
                print(f"\n❌ Erreur Gemini : {response.status_code}")
                print(response.text)
                return None
        except Exception as e:
            print(f"\n❌ Exception lors de l'appel à l'API Gemini : {str(e)}")
            return None
    
    # --- Génération et sauvegarde de la lettre de motivation ---
    def generer_et_sauvegarder_lettre(self, cv_dict: Dict, offre_dict: Dict, 
                                     infos_perso: Dict) -> Tuple[bool, Optional[str]]:
        """
        Génère une lettre de motivation et la sauvegarde dans un fichier Word
        
        Args:
            cv_dict: Dictionnaire contenant les informations du CV
            offre_dict: Dictionnaire contenant les informations de l'offre
            infos_perso: Informations supplémentaires pour personnaliser la lettre
            
        Returns:
            Tuple (succès, chemin_fichier) où:
            - succès est un booléen indiquant si la génération a réussi
            - chemin_fichier est le chemin vers le fichier généré ou None en cas d'échec
        """
        prompt_lettre = self.generer_prompt_lettre(cv_dict, offre_dict, infos_perso)
        lettre = self.interroger_gemini(prompt_lettre)
        
        if not lettre:
            return False, None
            
        try:
            # Création du document Word
            doc = Document()
            
            # En-tête avec coordonnées du candidat
            doc.add_paragraph(cv_dict['prenom_nom'])
            if 'adresse' in cv_dict:
                doc.add_paragraph(cv_dict['adresse'])
            doc.add_paragraph(cv_dict['email'])
            doc.add_paragraph(cv_dict['telephone'])
            if 'linkedin' in cv_dict:
                doc.add_paragraph(f"LinkedIn: {cv_dict['linkedin']}")
            
            # Espace
            doc.add_paragraph("")
            
            # Destinataire
            doc.add_paragraph(f"{offre_dict['entreprise']}")
            doc.add_paragraph(f"{offre_dict['lieu']}")
            
            # Date
            from datetime import date
            today = date.today()
            doc.add_paragraph(f"Le {today.strftime('%d/%m/%Y')}")
            
            # Objet
            doc.add_paragraph(f"Objet : Candidature au poste de {offre_dict['titre']}")
            
            # Corps de la lettre
            for ligne in lettre.split('\n'):
                if ligne.strip():
                    doc.add_paragraph(ligne)
            
            # Sauvegarde
            nom_fichier = f"{self.output_dir}/Lettre_{cv_dict['prenom_nom'].replace(' ', '_')}_{offre_dict['entreprise'].replace(' ', '_')}.docx"
            doc.save(nom_fichier)
            
            return True, nom_fichier
        except Exception as e:
            print(f"\n❌ Exception lors de la génération du document Word : {str(e)}")
            return False, None
    
    # --- Fonctions principales accessibles depuis l'extérieur ---
    def verifier_pertinence_cv(self, cv_dict: Dict, offre_dict: Dict) -> bool:
        """
        Vérifie si un CV est pertinent pour une offre d'emploi (>70% de compatibilité)
        
        Args:
            cv_dict: Dictionnaire contenant les informations du CV
            offre_dict: Dictionnaire contenant les informations de l'offre
            
        Returns:
            True si le CV est pertinent, False sinon
        """
        prompt = self.generer_prompt_pertinence(cv_dict, offre_dict)
        reponse = self.interroger_gemini(prompt)
        
        return reponse is not None and reponse.lower() == "oui"
    
    def generer_lettre_motivation(self, cv_dict: Dict, offre_dict: Dict, 
                                 infos_perso: Optional[Dict] = None) -> Tuple[bool, Optional[str], Optional[str]]:
        """
        Génère une lettre de motivation si le CV est pertinent pour l'offre
        
        Args:
            cv_dict: Dictionnaire contenant les informations du CV
            offre_dict: Dictionnaire contenant les informations de l'offre
            infos_perso: Informations supplémentaires pour personnaliser la lettre
            
        Returns:
            Tuple (succès, chemin_fichier, contenu_lettre) où:
            - succès est un booléen indiquant si la génération a réussi
            - chemin_fichier est le chemin vers le fichier généré ou None en cas d'échec
            - contenu_lettre est le contenu textuel de la lettre ou None en cas d'échec
        """
        if infos_perso is None:
            infos_perso = {"motivation": "", "lien_entreprise": "", "contraintes": ""}
            
        # Vérifier la pertinence du CV
        est_pertinent = self.verifier_pertinence_cv(cv_dict, offre_dict)
        
        if not est_pertinent:
            return False, None, None
        
        # Générer la lettre de motivation
        prompt_lettre = self.generer_prompt_lettre(cv_dict, offre_dict, infos_perso)
        contenu_lettre = self.interroger_gemini(prompt_lettre)
        
        if not contenu_lettre:
            return False, None, None
        
        # Sauvegarder la lettre dans un fichier Word
        succes, chemin_fichier = self.generer_et_sauvegarder_lettre(cv_dict, offre_dict, infos_perso)
        
        return succes, chemin_fichier, contenu_lettre

# Fonction pour être utilisée depuis bot.py
def integrer_module_groupe5(api_key, output_dir="lettres"):
    """
    Crée et renvoie une instance du module du groupe 5
    
    Args:
        api_key: Clé API Google Gemini
        output_dir: Répertoire où seront sauvegardées les lettres générées
        
    Returns:
        Une instance de ModuleGroupe5
    """
    return ModuleGroupe5(api_key, output_dir)